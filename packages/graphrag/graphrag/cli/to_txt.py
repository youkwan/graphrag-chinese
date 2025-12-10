import shutil
import tempfile
from pathlib import Path

import typer


def _convert_doc_to_docx(doc_path: Path) -> Path:
    """Convert a .doc file to a temporary .docx using Microsoft Word."""
    try:
        import win32com.client  # type: ignore
    except ImportError as exc:  # pragma: no cover - depends on local install
        msg = "Processing .doc requires Microsoft Word and pywin32."
        raise RuntimeError(msg) from exc

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = None
    temp_dir = Path(tempfile.mkdtemp(prefix="graphrag_doc_"))
    try:
        doc = word.Documents.Open(str(doc_path))
        out_path = temp_dir / f"{doc_path.stem}.docx"
        doc.SaveAs(str(out_path), FileFormat=16)  # 16 == wdFormatXMLDocument
    except Exception as exc:  # pragma: no cover - depends on local install
        msg = "Failed to convert .doc to .docx; ensure Microsoft Word is installed and activated."
        raise RuntimeError(msg) from exc
    else:
        return out_path
    finally:
        if doc is not None:
            doc.Close(False)
        word.Quit()


def _extract_docx_text(docx_path: Path) -> str:
    """Load a .docx via python-docx and return plain text."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - depends on local install
        msg = "python-docx is required; please install python-docx."
        raise RuntimeError(msg) from exc

    doc = Document(str(docx_path))
    lines: list[str] = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)

    return "\n".join(lines)


def _load_word_text(doc_path: Path) -> str:
    """Convert a Word file to plain text."""
    suffix = doc_path.suffix.lower()
    target_path = doc_path

    if suffix == ".doc":
        temp_docx: Path | None = None
        try:
            temp_docx = _convert_doc_to_docx(doc_path)
            target_path = temp_docx
            return _extract_docx_text(target_path)
        finally:
            if temp_docx and temp_docx.exists():
                shutil.rmtree(temp_docx.parent, ignore_errors=True)

    if suffix == ".docx":
        return _extract_docx_text(target_path)

    msg = "Only .doc or .docx files are supported."
    raise RuntimeError(msg)


def convert_word_docs_to_txt(
    source_dir: Path,
    output_dir: Path,
    overwrite: bool = False,
) -> None:
    """Convert all Word files under a directory tree into txt."""
    if not source_dir.exists() or not source_dir.is_dir():
        msg = f"Source directory not found: {source_dir}"
        raise FileNotFoundError(msg)

    output_dir.mkdir(parents=True, exist_ok=True)

    doc_files = sorted(source_dir.rglob("*.doc")) + sorted(source_dir.rglob("*.docx"))
    if not doc_files:
        typer.echo("No .doc or .docx files found.")
        return

    converted = 0
    skipped = 0
    for idx, doc_path in enumerate(doc_files, start=1):
        target = output_dir / f"data_{idx:07d}.txt"
        if target.exists() and not overwrite:
            skipped += 1
            continue

        text_content = _load_word_text(doc_path)
        target.write_text(text_content, encoding="utf-8", errors="ignore")
        converted += 1

    typer.echo(f"Wrote {converted} files to {output_dir.resolve()}, skipped {skipped} existing files.")


def to_txt_cli(
    source: Path = typer.Argument(
        ...,
        help="Source folder; all .doc/.docx under it will be processed recursively.",
        exists=True,
        file_okay=False,
        dir_okay=True,
        readable=True,
        resolve_path=True,
    ),
    output: Path = typer.Option(
        ...,
        "--output",
        "-o",
        help="Output folder where converted txt files are written. (Required)",
        dir_okay=True,
        writable=True,
        resolve_path=True,
    ),
    overwrite: bool = typer.Option(
        False,
        "--overwrite/--no-overwrite",
        help="Overwrite existing txt files if present.",
    ),
) -> None:
    """Typer command: convert Word files to txt and gather them into one folder."""
    convert_word_docs_to_txt(source_dir=source, output_dir=output, overwrite=overwrite)
